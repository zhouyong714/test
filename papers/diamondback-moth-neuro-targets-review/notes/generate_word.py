#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generate Word document from the abstract content
"""

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Error: python-docx is not installed.")
    print("Please install it using: pip install python-docx")
    exit(1)

def create_abstract_document():
    doc = Document()

    # Title
    title = doc.add_heading('小菜蛾神经系统相关药理靶标概述：作用机制、抗性位点与靶标发现前景', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Abstract heading
    doc.add_heading('摘要', level=2)

    # Abstract paragraphs
    paragraphs = [
        "小菜蛾（Plutella xylostella）是全球十字花科作物上最具代表性的高抗性害虫之一，其对几乎所有主流杀虫剂类别均已产生抗性，使其成为研究杀虫剂靶标与抗性共演化的理想模型。神经系统作用剂长期构成小菜蛾化学防治的核心，现有研究已在乙酰胆碱酯酶（AChE）、烟碱型乙酰胆碱受体（nAChR）、GABA受体（RDL）、谷氨酸门控氯离子通道（GluCl）、电压门控钠通道（Na_v）和兰尼碱受体（RyR）等靶标上积累了大量数据。然而，这些知识体系仍主要沿药剂类别分散组织，缺少一个以"神经系统药理靶标"为中心的统一框架，将作用机制、靶标位点、抗性突变、交叉抗性风险和新靶标发现整合讨论。",

        "本综述将小菜蛾神经系统相关靶标重组为四个功能模块：（1）胆碱能兴奋性通路（AChE与nAChR），（2）抑制性氯离子通路（RDL/GABA受体与GluCl），（3）膜兴奋性与胞内钙释放通路（Na_v与RyR），（4）新兴生物胺受体靶标（酪胺受体与章鱼胺受体）。针对每个靶标模块，本文系统梳理了靶标生物学基础、代表性药剂及其作用位点、小菜蛾中的直接功能证据、关键抗性突变的分子机制，以及这些发现对抗性监测与综合治理的意义。",

        "在胆碱能通路中，nAChR α6亚基的错误剪接、提前终止密码子和跨膜区三氨基酸缺失为多杀菌素类抗性提供了完整的因果证据链。在抑制性氯离子通路中，GluCl的A309V、G315E、V263I突变及外显子9剪接变体构成了阿维菌素类抗性的核心机制，而小菜蛾特有的两个RDL同源基因为理解非芳基吡唑类与新型间二酰胺类药剂的作用差异提供了独特视角。在膜兴奋性与钙释放通路中，RyR的G4946E、I4790M和I4790K突变已通过CRISPR基因编辑获得功能验证，使小菜蛾成为双酰胺类抗性研究中证据最完整的昆虫物种；Na_v的F1845Y和V1848I突变则将传统拟除虫菊酯抗性研究与新型钠通道阻断剂（SCBI）抗性机制连接起来。在新兴靶标方面，酪胺受体1（TAR1）和章鱼胺受体β3（Octβ3）的药理学特性、RNAi功能验证和CRISPR敲除表型研究表明，生物胺GPCR系统具备"可药理学操控、可功能验证、可用于新药设计"的三重潜力。",

        "本文构建了小菜蛾神经靶标-药剂类别-关键抗性位点的统一对照图谱，提出了基于证据强度的靶标优先级评估框架，并整合了适用于分子监测的标记物组合（Pxace1替换、Pxα6错误剪接与缺失、PxGluCl A309V/G315E/V263I、PxNa_v F1845Y/V1848I、PxRyR G4946E/I4790M/I4790K）。同时，本文强调靶标位点突变并非抗性的唯一机制，代谢解毒、转运蛋白上调和共生菌介导的抗性在双酰胺类、阿维菌素类和拟除虫菊酯类体系中均有重要贡献，提示抗性监测需要整合靶标与非靶标机制。",

        "展望未来，本文提出了面向新靶标发现的方法学路线图，包括异源表达与电生理验证、CRISPR介导的功能遗传学、受体-配体对接与分子动力学模拟、单细胞表达图谱与神经解剖定位，以及群体监测与高通量基因分型技术的整合。短期目标（6-18个月）应聚焦于建立标准化分子标记组、汇总全球代表群体的等位基因频率、建立TAR1/Octβ3的稳定异源表达与药理筛选体系；中期目标（2-5年）应致力于将结构生物学与田间抗性预测相连接，针对变构位点或GPCR设计下一代选择性化学先导，并建立"靶标图谱-抗性位点数据库-抗性治理建议"的一体化平台。"
    ]

    for para_text in paragraphs:
        para = doc.add_paragraph(para_text)
        para.paragraph_format.first_line_indent = Inches(0.5)
        para.paragraph_format.line_spacing = 1.5
        for run in para.runs:
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'

    # Keywords
    keywords_para = doc.add_paragraph()
    keywords_run = keywords_para.add_run('关键词：')
    keywords_run.bold = True
    keywords_run.font.size = Pt(12)
    keywords_text = keywords_para.add_run('小菜蛾；神经系统；药理靶标；抗性机制；靶标位点突变；分子监测；新靶标发现')
    keywords_text.font.size = Pt(12)

    # References section
    doc.add_page_break()
    doc.add_heading('参考文献', level=2)

    references = [
        "Sparks TC, Crossthwaite AJ, Nauen R, et al. Insecticides, biologics and nematicides: Updates to IRAC's mode of action classification - a tool for resistance management. Pestic Biochem Physiol. 2020;167:104587.",

        "Furlong MJ, Wright DJ, Dosdall LM. Diamondback moth ecology and management: problems, progress, and prospects. Annu Rev Entomol. 2013;58:517-541.",

        "You M, Yue Z, He W, et al. A heterozygous moth genome provides insights into herbivory and detoxification. Nat Genet. 2013;45:220-225.",

        "Baxter SW, Chen M, Dawson A, et al. Mis-spliced transcripts of nicotinic acetylcholine receptor alpha6 are associated with field evolved spinosad resistance in Plutella xylostella. PLoS Genet. 2010;6(1):e1000802.",

        "Wang X, Wang R, Yang Y, et al. A point mutation in the glutamate-gated chloride channel of Plutella xylostella is associated with resistance to abamectin. Insect Mol Biol. 2016;25(2):116-125.",

        "Troczka B, Zimmer CT, Elias J, et al. Resistance to diamide insecticides in diamondback moth is associated with a mutation in the membrane-spanning domain of the ryanodine receptor. Insect Biochem Mol Biol. 2012;42(11):873-880.",

        "Wang XL, Su W, Zhang JH, et al. Two novel sodium channel mutations associated with resistance to indoxacarb and metaflumizone in the diamondback moth, Plutella xylostella. Insect Sci. 2016;23(1):50-58.",

        "Ma H, Huang Q, Lai X, et al. Pharmacological properties of the type 1 tyramine receptor in the diamondback moth, Plutella xylostella. Int J Mol Sci. 2019;20(12):2953."
    ]

    for i, ref in enumerate(references, 1):
        ref_para = doc.add_paragraph(f'[{i}] {ref}')
        ref_para.paragraph_format.line_spacing = 1.5
        for run in ref_para.runs:
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'

    # Save document
    output_path = '小菜蛾神经系统药理靶标综述_摘要.docx'
    doc.save(output_path)
    print(f"Word document created successfully: {output_path}")

if __name__ == '__main__':
    create_abstract_document()
