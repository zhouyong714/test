from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

HTML_PATH = r"D:\papers\codex\papers\diamondback-moth-neuro-targets-review\notes\全文合并.html"
DOCX_PATH = r"D:\papers\codex\papers\diamondback-moth-neuro-targets-review\notes\全文合并.docx"

FONT_EN = "Times New Roman"
FONT_ZH = "SimSun"


def set_run_font(run, size_pt=12, bold=False, italic=False):
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size_pt)
    run.font.name = FONT_EN
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_ZH)


def set_para_font(para, size_pt=12, bold=False, indent_first=True):
    pf = para.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    if indent_first:
        pf.first_line_indent = Cm(0.74)  # ~2 Chinese chars
    for run in para.runs:
        run.font.size = Pt(size_pt)
        run.font.name = FONT_EN
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_ZH)
        run.bold = bold


def add_inline_content(para, element, base_size=12, bold=False, italic=False):
    """Recursively add inline content from an element into a paragraph."""
    for child in element.children:
        if isinstance(child, str):
            text = child
            if text:
                run = para.add_run(text)
                set_run_font(run, base_size, bold=bold, italic=italic)
        elif child.name == "sup":
            run = para.add_run(child.get_text())
            run.font.size = Pt(base_size * 0.75)
            run.font.name = FONT_EN
            run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_ZH)
            run.font.superscript = True
        elif child.name in ("em", "i"):
            add_inline_content(para, child, base_size, bold=bold, italic=True)
        elif child.name in ("strong", "b"):
            add_inline_content(para, child, base_size, bold=True, italic=italic)
        elif child.name == "sub":
            run = para.add_run(child.get_text())
            run.font.size = Pt(base_size * 0.75)
            run.font.name = FONT_EN
            run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_ZH)
            run.font.subscript = True
        else:
            add_inline_content(para, child, base_size, bold=bold, italic=italic)


def add_table(doc, table_tag):
    rows = table_tag.find_all("tr")
    if not rows:
        return
    col_count = max(len(r.find_all(["td", "th"])) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=col_count)
    tbl.style = "Table Grid"
    for i, row in enumerate(rows):
        cells = row.find_all(["td", "th"])
        for j, cell in enumerate(cells):
            if j >= col_count:
                break
            tc = tbl.rows[i].cells[j]
            para = tc.paragraphs[0]
            is_header = cell.name == "th"
            add_inline_content(para, cell, base_size=10, bold=is_header)
            for run in para.runs:
                run.font.size = Pt(10)
                run.font.name = FONT_EN
                run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_ZH)


def process_body(doc, body):
    for el in body.children:
        if not hasattr(el, "name") or el.name is None:
            continue

        if el.name == "h1":
            p = doc.add_heading(level=1)
            p.clear()
            add_inline_content(p, el, base_size=16, bold=True)
            p.paragraph_format.first_line_indent = Pt(0)
            p.alignment = 1  # center

        elif el.name == "h2":
            p = doc.add_heading(level=2)
            p.clear()
            add_inline_content(p, el, base_size=14, bold=True)
            p.paragraph_format.first_line_indent = Pt(0)

        elif el.name == "h3":
            p = doc.add_heading(level=3)
            p.clear()
            add_inline_content(p, el, base_size=12, bold=True)
            p.paragraph_format.first_line_indent = Pt(0)

        elif el.name == "p":
            cls = el.get("class", [])
            is_ref = "ref-item" in cls
            is_kw = "keywords" in cls
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.space_after = Pt(4)
            if is_ref:
                pf.first_line_indent = Cm(-0.74)
                pf.left_indent = Cm(0.74)
            elif is_kw:
                pf.first_line_indent = Pt(0)
            else:
                pf.first_line_indent = Cm(0.74)
            add_inline_content(p, el, base_size=10 if is_ref else 12)

        elif el.name == "table":
            caption = el.find("caption")
            if caption:
                cp = doc.add_paragraph()
                cp.paragraph_format.first_line_indent = Pt(0)
                run = cp.add_run(caption.get_text())
                run.italic = True
                run.font.size = Pt(10)
                run.font.name = FONT_EN
                run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_ZH)
            add_table(doc, el)

        elif el.name == "hr":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)

        elif el.name == "div":
            process_body(doc, el)


def main():
    with open(HTML_PATH, encoding="utf-8") as f:
        soup = BeautifulSoup(f, "lxml")

    doc = Document()
    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    # Default style
    style = doc.styles["Normal"]
    style.font.name = FONT_EN
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_ZH)

    body = soup.find("body")
    process_body(doc, body)

    doc.save(DOCX_PATH)
    print(f"Saved: {DOCX_PATH}")


if __name__ == "__main__":
    main()
